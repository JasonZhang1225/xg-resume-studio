"""滴鱼简历助手 XG Resume Studio —— 本地运行、数据不出电脑的个人简历管理系统。"""
import io
import json
import os
import re
import secrets
import shutil
import socket
import sqlite3
import sys
import uuid
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, File, Request, UploadFile, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from jinja2 import Environment, select_autoescape

import database as db
import cert_organizer
import extractor
import parsers

# PyInstaller 打包态：模板/静态资源在解包目录（只读），用户数据固定在 exe 旁边（可写）
FROZEN = getattr(sys, "frozen", False)
BASE_DIR = Path(getattr(sys, "_MEIPASS", "")) if FROZEN else Path(__file__).resolve().parent
APP_ROOT = Path(sys.executable).resolve().parent if FROZEN else BASE_DIR
# 用户运行时数据统一放在 data/ 下（数据库/备份/上传件），打包分享时整体排除
RUNTIME_DIR = APP_ROOT / "data"
UPLOAD_DIR = RUNTIME_DIR / "uploads"
PHOTO_DIR = UPLOAD_DIR / "photos"
ATTACH_DIR = UPLOAD_DIR / "attachments"
CERT_DIR = UPLOAD_DIR / "certificates"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
PHOTO_DIR.mkdir(exist_ok=True)
ATTACH_DIR.mkdir(exist_ok=True)
CERT_DIR.mkdir(exist_ok=True)

ALLOWED_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tif", ".tiff", ".pdf", ".docx"}
ATTACH_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".webp", ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".zip"}
CERT_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tif", ".tiff", ".pdf"}

app = FastAPI(title="滴鱼简历助手 XG Resume Studio")
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")
app.mount("/uploads", StaticFiles(directory=str(UPLOAD_DIR)), name="uploads")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

db.init_db()
# 启动时自动备份数据库，防止误操作丢数据（备份在 data_backups/）
try:
    db.backup_db()
except Exception:
    pass
# 后台预热 OCR 模型，让首次上传识别更快（未装 OCR 时静默跳过）
parsers.warmup_ocr()

# 老实例兼容：升级前已有真实数据（资料/条目）的库视为已完成初始化，不再弹向导
if db.get_setting("initialized") != "1":
    _prof = db.get_row("profile", 1) or {}
    _has_content = bool((_prof.get("name") or "").strip()) or any(
        db.get_rows_where(t, "1=1", (), "id")
        for t in ("awards", "positions", "education", "papers", "projects", "files"))
    if _has_content:
        db.set_setting("initialized", "1")


def _needs_setup() -> bool:
    return db.get_setting("initialized") != "1"


# ---------- 局域网配对码（仅 RESUME_LAN=1 时生效） ----------

_PAIR_ALPHABET = "ABCDEFGHJKLMNPQRSTUVWXYZ23456789"  # 去掉易混淆的 0/O/1/I


def _pair_code() -> str:
    code = db.get_setting("pair_code")
    if not code:
        code = "".join(secrets.choice(_PAIR_ALPHABET) for _ in range(6))
        db.set_setting("pair_code", code)
    return code


def _is_loopback(host: str) -> bool:
    return host in ("127.0.0.1", "::1", "localhost")


@app.middleware("http")
async def _entry_gate(request, call_next):
    """入口门卫：① 未完成向导 → 重定向 /setup；
    ② 局域网模式下，非本机设备必须持有效配对 Cookie，否则引导到 /pair 输码。"""
    if _needs_setup():
        p = request.url.path
        allowed = ("/setup", "/static/", "/api/setup", "/api/ai/test", "/api/assistant/avatar")
        if not any(p == a or p.startswith(a) for a in allowed):
            if p.startswith("/api/") or request.method != "GET":
                return JSONResponse({"detail": "请先完成初始化向导", "need_setup": True}, status_code=403)
            return RedirectResponse("/setup")

    if os.environ.get("RESUME_LAN") == "1":
        client_ip = request.client.host if request.client else ""
        if not _is_loopback(client_ip):
            got = request.cookies.get("pair") or ""
            if not secrets.compare_digest(got, _pair_code()):
                p = request.url.path
                if p == "/pair" or p.startswith("/api/pair"):
                    pass  # 配对页与配对接口本身放行
                elif p.startswith("/api/") or request.method != "GET":
                    return JSONResponse({"detail": "请先输入配对码", "need_pair": True}, status_code=403)
                else:
                    from urllib.parse import quote
                    return RedirectResponse("/pair?next=" + quote(p))

    return await call_next(request)


@app.middleware("http")
async def _csrf_guard(request, call_next):
    """CSRF 防护：浏览器跨站发起的 POST/PUT/DELETE 会携带外部 Origin，
    与请求 Host 不一致即拒绝。恶意网页因此无法驱动你本机的服务删数据。
    （无 Origin 的请求来自 curl/同源表单等，不受影响。）"""
    if request.method in ("POST", "PUT", "DELETE", "PATCH"):
        origin = request.headers.get("origin")
        if origin:
            from urllib.parse import urlsplit
            host = request.headers.get("host", "")
            if urlsplit(origin).netloc != host:
                return JSONResponse({"detail": "跨站请求被拒绝"}, status_code=403)
    return await call_next(request)


AWARD_COLS = ["title", "level", "category", "date", "organizer", "description", "source_file"]
POSITION_COLS = ["title", "org", "start", "end", "description", "source_file", "ongoing"]
EDU_COLS = ["school", "degree", "major", "start", "end", "notes", "ongoing"]
PAPER_COLS = ["title", "authors", "venue", "year", "volume", "issue", "pages", "paper_type", "notes"]
PROJECT_COLS = ["name", "role", "start", "end", "description", "ongoing"]
PROFILE_COLS = ["name", "gender", "birth_date", "phone", "email", "address", "hometown", "summary",
                "skills", "languages", "photo_path"]

ITEM_TABLES = {
    "award": ("awards", AWARD_COLS),
    "position": ("positions", POSITION_COLS),
}


# ---------- 多账户 ----------

def _uid(request: Request) -> int:
    """当前账户：手机令牌 > Cookie > 默认第一个账户。"""
    tok = request.headers.get("x-device-token", "")
    if tok:
        for u in db.all_rows("users", "id ASC"):
            if tok == db.get_setting(f"mobile_token:{u['id']}"):
                return u["id"]
    try:
        uid = int(request.cookies.get("uid", "1"))
    except (TypeError, ValueError):
        uid = 1
    user = db.get_row("users", uid)
    return user["id"] if user else 1


def _base_ctx(request: Request, **extra):
    uid = _uid(request)
    ctx = {
        "active": "",
        "assistant_avatar": db.get_setting("assistant_avatar"),
        "cur_uid": uid,
        "cur_user": db.get_row("users", uid) or {"id": uid, "name": "?", "avatar": ""},
        "users": db.all_rows("users", "id ASC"),
        "lan_mode": os.environ.get("RESUME_LAN") == "1",
        "is_remote": not _is_loopback(request.client.host if request.client else "127.0.0.1"),
    }
    ctx.update(extra)
    return ctx


def _pick(data, cols):
    out = {}
    for c in cols:
        v = data.get(c, "")
        out[c] = "" if v is None else v
    return out


def _now():
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def _remove_local_file(path):
    if path and os.path.exists(path):
        try:
            os.unlink(path)
        except OSError:
            pass


def _remove_upload_file(public_url):
    """删除 /uploads/... 对应的本机旧文件（仅限 uploads 目录内，防路径逃逸）。"""
    if not public_url or not public_url.startswith("/uploads/"):
        return
    root = UPLOAD_DIR.resolve()
    target = (UPLOAD_DIR / public_url[len("/uploads/"):]).resolve()
    try:
        target.relative_to(root)
    except ValueError:
        return
    if target.is_file():
        try:
            target.unlink()
        except OSError:
            pass


def _to_flag(value):
    return 1 if str(value).strip().lower() in ("1", "true", "on", "yes") else 0


def _mark_ongoing(rows):
    """进行中的学历/任职/项目：结束时间统一显示为「至今」，排序也天然置顶。"""
    for r in rows:
        if r.get("ongoing"):
            r["end"] = "至今"
    return rows


def _date_key(value):
    """把「2024年10月 / 2024-6 / 2023.09 / 2024年 / 至今」等归一化成可排序的 (年, 月)。"""
    if not value:
        return (0, 0)
    text = str(value)
    if "至今" in text or "现在" in text:
        return (9999, 12)
    m = re.search(r"(?<!\d)(\d{4})\s*年\s*(\d{1,2})\s*月", text)
    if not m:
        m = re.search(r"(?<!\d)(\d{4})\s*[./\-]\s*(\d{1,2})(?!\d)", text)
    if m and 1 <= int(m.group(2)) <= 12:
        return (int(m.group(1)), int(m.group(2)))
    m = re.search(r"(?<!\d)(\d{4})\s*年", text) or re.search(r"(?<!\d)(\d{4})(?!\d)", text)
    if m:
        return (int(m.group(1)), 0)
    return (0, 0)


def paper_citation(a: dict) -> str:
    """把论文条目合成 GB/T 7714 风格引用串：
    作者. 标题[J]. 期刊, 年份, 卷(期): 起止页码.
    类型含「会议」→[C]，含「学位」→[D]，其余默认 [J]。"""
    authors = str(a.get("authors") or "").strip().rstrip(".")
    title = str(a.get("title") or "").strip()
    venue = str(a.get("venue") or "").strip()
    year = str(a.get("year") or "").strip()
    ptype = str(a.get("paper_type") or "").strip()
    vol = str(a.get("volume") or "").strip()
    issue = str(a.get("issue") or "").strip()
    pages = str(a.get("pages") or "").strip()
    marker = "[C]" if "会议" in ptype else ("[D]" if "学位" in ptype else "[J]")
    head = (authors + ". " if authors else "") + (title + marker + ". " if title else "")
    rest = ", ".join(x for x in (venue, year) if x)
    if vol or issue or pages:
        vi = f"{vol}({issue})" if vol and issue else (vol or (f"({issue})" if issue else ""))
        rest = ", ".join(x for x in (rest, f"{vi}:{pages}" if vi and pages else (vi or pages)) if x)
    out = (head + rest).strip()
    return out + "." if out and not out.endswith(".") else out


# ---------- 简历模板 ----------

TEMPLATE_DIR = BASE_DIR / "resume_templates"
DEFAULT_TEMPLATE_ORDER = ["classic", "scholar", "career", "minimal"]
SECTION_KEYS = ["summary", "education", "papers", "projects", "awards", "positions", "skills"]

_tpl_meta_cache = None


def _tpl_meta():
    global _tpl_meta_cache
    if _tpl_meta_cache is None:
        p = TEMPLATE_DIR / "_meta.json"
        try:
            _tpl_meta_cache = json.loads(p.read_text(encoding="utf-8")) if p.exists() else {}
        except Exception:
            _tpl_meta_cache = {}
    return _tpl_meta_cache


def seed_builtin_templates():
    if not TEMPLATE_DIR.exists():
        return
    for f in sorted(TEMPLATE_DIR.glob("*.html")):
        html = f.read_text(encoding="utf-8")
        meta = _tpl_meta().get(f.stem, {})
        title = meta.get("title", f.stem)
        conn = db.get_conn()
        # upsert：内置模板每次启动都用文件内容刷新，改 resume_templates/*.html 立即生效
        conn.execute(
            """INSERT INTO resume_templates (name, title, html, builtin, created_at)
               VALUES (?, ?, ?, 1, ?)
               ON CONFLICT(name) DO UPDATE SET html = excluded.html, title = excluded.title, builtin = 1""",
            (f.stem, title, html, _now()),
        )
        conn.commit()
        conn.close()
    # 清理已被替换/删除的旧内置模板行（自定义模板不受影响）
    stems = {f.stem for f in TEMPLATE_DIR.glob("*.html")}
    conn = db.get_conn()
    marks = ",".join("?" * len(stems))
    conn.execute(f"DELETE FROM resume_templates WHERE builtin = 1 AND name NOT IN ({marks})", tuple(stems))
    conn.commit()
    conn.close()


seed_builtin_templates()

_custom_env = Environment(autoescape=select_autoescape(default_for_string=True, default=True))


def _base_css():
    p = TEMPLATE_DIR / "_base.css"
    return p.read_text(encoding="utf-8") if p.exists() else ""


def render_resume_template(name, ctx):
    row = db.get_row_by("resume_templates", "name", name)
    if not row:
        raise HTTPException(404, f"模板 {name} 不存在")
    base = f"<style>\n{_base_css()}\n</style>\n" if _base_css() else ""
    return base + _custom_env.from_string(row["html"]).render(**ctx)


def resume_context(uid: int):
    profile = db.get_row("profile", uid) or {}
    profile["skills"] = db.json_list(profile.get("skills"))
    profile["languages"] = db.json_list(profile.get("languages"))
    counts = db.attachment_counts()
    awards = db.get_rows_where("awards", "user_id=?", (uid,))
    positions = db.get_rows_where("positions", "user_id=?", (uid,))
    for a in awards:
        a["attachments"] = db.attachments_for("award", a["id"])
    for p in positions:
        p["attachments"] = db.attachments_for("position", p["id"])
    papers = db.get_rows_where("papers", "user_id=?", (uid,))
    for a in papers:
        a["citation"] = paper_citation(a)
    return {
        "profile": profile,
        "education": sorted(_mark_ongoing(db.get_rows_where("education", "user_id=?", (uid,))),
                            key=lambda r: _date_key(r.get("end")), reverse=True),
        "papers": papers,
        "projects": _mark_ongoing(db.get_rows_where("projects", "user_id=?", (uid,))),
        "awards": sorted(awards, key=lambda r: _date_key(r.get("date")), reverse=True),
        "positions": _mark_ongoing(sorted(positions, key=lambda r: _date_key(r.get("start")), reverse=True)),
        "attach_counts": counts,
    }


# ---------- 页面 ----------

@app.get("/", response_class=HTMLResponse)
def index(request: Request):
    uid = _uid(request)
    return templates.TemplateResponse(request, "index.html", _base_ctx(request,
        active="home",
        files=db.get_rows_where("files", "user_id=?", (uid,), "id DESC"),
        award_count=len(db.get_rows_where("awards", "user_id=?", (uid,))),
        position_count=len(db.get_rows_where("positions", "user_id=?", (uid,))),
        cert_count=len(db.get_rows_where("certificates", "user_id=?", (uid,))),
        has_ai=bool(db.get_setting("api_key")),
        lan_mode=os.environ.get("RESUME_LAN") == "1",
    ))


@app.get("/items", response_class=HTMLResponse)
def items_page(request: Request):
    uid = _uid(request)
    return templates.TemplateResponse(request, "items.html", _base_ctx(request,
        active="items",
        awards=db.get_rows_where("awards", "user_id=?", (uid,)),
        positions=_mark_ongoing(db.get_rows_where("positions", "user_id=?", (uid,))),
        attach_counts=db.attachment_counts(),
    ))


@app.get("/profile", response_class=HTMLResponse)
def profile_page(request: Request):
    uid = _uid(request)
    profile = db.get_row("profile", uid) or {}
    profile["skills"] = db.json_list(profile.get("skills"))
    profile["languages"] = db.json_list(profile.get("languages"))
    return templates.TemplateResponse(request, "profile.html", _base_ctx(request,
        active="profile",
        profile=profile,
        education=db.get_rows_where("education", "user_id=?", (uid,)),
        papers=db.get_rows_where("papers", "user_id=?", (uid,)),
        projects=db.get_rows_where("projects", "user_id=?", (uid,)),
    ))


@app.get("/chat", response_class=HTMLResponse)
def chat_page(request: Request):
    return templates.TemplateResponse(request, "chat.html", _base_ctx(request,
        active="chat",
        has_ai=bool(db.get_setting("api_key")),
    ))


@app.get("/resume", response_class=HTMLResponse)
def resume_page(request: Request, template: str = "", accent: str = "", density: str = "standard"):
    uid = _uid(request)
    avail = {r["name"]: r for r in db.all_rows("resume_templates", "builtin DESC, id ASC")}
    if template not in avail:
        # 未指定/已失效（如旧模板被替换）时回退到默认顺序里的第一个
        template = next((n for n in DEFAULT_TEMPLATE_ORDER if n in avail), next(iter(avail), ""))
    if not template:
        raise HTTPException(404, "暂无可用模板，请重启服务以载入内置模板")
    meta = _tpl_meta().get(template, {})
    accent = accent if re.fullmatch(r"#[0-9a-fA-F]{6}", accent or "") else meta.get("accent", "#0f766e")
    if density not in ("compact", "standard", "relaxed"):
        density = "standard"

    ctx = resume_context(uid)
    ctx["templates"] = db.all_rows("resume_templates", "builtin DESC, id ASC")
    ctx["current_template"] = template
    ctx["body_html"] = render_resume_template(template, ctx)
    ctx["accent"] = accent
    ctx["density"] = density
    ctx["tmeta"] = _tpl_meta()
    return templates.TemplateResponse(request, "resume_shell.html", _base_ctx(request, **ctx))


# ---------- 简历模块布局 ----------

def _normalize_layout(raw):
    try:
        data = json.loads(raw) if raw else {}
    except Exception:
        data = {}
    order = [k for k in (data.get("order") or []) if k in SECTION_KEYS]
    hidden = [k for k in (data.get("hidden") or []) if k in SECTION_KEYS]
    order += [k for k in SECTION_KEYS if k not in order]
    hidden = [k for k in hidden if k in order]
    return {"order": order, "hidden": hidden}


@app.get("/api/layout")
def get_layout(request: Request):
    return _normalize_layout(db.get_setting(f"resume_layout:{_uid(request)}"))


@app.put("/api/layout")
def put_layout(request: Request, payload: dict):
    merged = _normalize_layout(json.dumps({
        "order": payload.get("order", []),
        "hidden": payload.get("hidden", []),
    }, ensure_ascii=False))
    db.set_setting(f"resume_layout:{_uid(request)}", json.dumps(merged, ensure_ascii=False))
    return {"ok": True, **merged}


# ---------- 上传与解析 ----------
# 注意：解析/OCR 是重活，路由用同步 def 让 FastAPI 自动放进线程池，避免阻塞事件循环。

@app.post("/api/upload")
def upload_file(request: Request, file: UploadFile = File(...)):
    uid = _uid(request)
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTS:
        raise HTTPException(400, f"不支持的格式：{ext or '未知'}，支持 jpg/png/pdf/docx")
    stored = UPLOAD_DIR / f"{uuid.uuid4().hex}{ext}"
    with open(stored, "wb") as f:
        f.write(file.file.read())
    raw_text = parsers.extract_text(str(stored), ext)
    cand = extractor.extract_candidates(raw_text)
    file_id = db.insert_row("files", {
        "user_id": uid,
        "filename": file.filename or "",
        "stored_path": str(stored),
        "file_type": ext,
        "status": "parsed",
        "raw_text": raw_text,
        "created_at": _now(),
    })
    return {
        "file_id": file_id,
        "filename": file.filename,
        "file_type": ext,
        "raw_preview": raw_text[:2000],
        "awards": cand["awards"],
        "positions": cand["positions"],
    }


@app.get("/api/files")
def list_files(request: Request):
    return db.get_rows_where("files", "user_id=?", (_uid(request),), "id DESC")


@app.get("/api/files/{file_id}")
def get_file(request: Request, file_id: int):
    row = db.get_row("files", file_id)
    if not row or row["user_id"] != _uid(request):
        raise HTTPException(404, "文件不存在")
    return row


@app.delete("/api/files/{file_id}")
def delete_file(request: Request, file_id: int):
    row = db.get_row("files", file_id)
    if not row or row["user_id"] != _uid(request):
        raise HTTPException(404, "文件不存在")
    _remove_local_file(row["stored_path"])
    db.delete_row("files", file_id)
    return {"ok": True}


@app.post("/api/files/{file_id}/rotate")
def rotate_file(request: Request, file_id: int, payload: dict):
    """把上传的证书图片旋转后重新 OCR（自动方向校正救不回来时的手动兜底）。"""
    uid = _uid(request)
    row = db.get_row("files", file_id)
    if not row or row["user_id"] != uid:
        raise HTTPException(404, "文件不存在")
    if (row.get("file_type") or "") not in parsers.IMAGE_EXTS:
        raise HTTPException(400, "仅图片文件支持旋转重试")
    try:
        angle = ((int(payload.get("angle", 0)) % 360) + 360) % 360
    except (TypeError, ValueError):
        raise HTTPException(400, "angle 必须是 90/180/270")
    if angle not in (90, 180, 270):
        raise HTTPException(400, "angle 必须是 90/180/270")
    path = row["stored_path"]
    if not path or not os.path.exists(path):
        raise HTTPException(404, "原始文件已丢失，无法旋转")
    if not parsers.rotate_image_file(path, angle):
        raise HTTPException(500, "旋转失败，请确认已安装 Pillow")
    raw_text = parsers.extract_text(path, row["file_type"])
    cand = extractor.extract_candidates(raw_text)
    db.update_row("files", file_id, {"raw_text": raw_text})
    return {
        "file_id": file_id,
        "filename": row["filename"],
        "file_type": row["file_type"],
        "raw_preview": raw_text[:2000],
        "awards": cand["awards"],
        "positions": cand["positions"],
    }


# ---------- 获奖 / 任职 ----------

@app.get("/api/items")
def list_items(request: Request):
    uid = _uid(request)
    counts = db.attachment_counts()
    awards = db.get_rows_where("awards", "user_id=?", (uid,))
    positions = db.get_rows_where("positions", "user_id=?", (uid,))
    for a in awards:
        a["attach_count"] = counts.get(("award", a["id"]), 0)
    for p in positions:
        p["attach_count"] = counts.get(("position", p["id"]), 0)
    return {"awards": awards, "positions": positions}


@app.post("/api/items")
def create_item(request: Request, payload: dict):
    uid = _uid(request)
    item_type = payload.get("item_type", "")
    table, cols = ITEM_TABLES.get(item_type, (None, None))
    if not table:
        raise HTTPException(400, "item_type 必须是 award 或 position")
    data = _pick(payload, cols)
    if "ongoing" in data:
        data["ongoing"] = _to_flag(data["ongoing"])
    data["user_id"] = uid
    data["sort_order"] = db.next_sort_order(table)
    new_id = db.insert_row(table, data)
    return {"id": new_id}


# 批量导入查重的键：标题 + 时间都相同视为同一份材料，避免重复上传产生重复条目
_DEDUP_DATE_COL = {"awards": "date", "positions": "start"}


@app.post("/api/items/batch")
def batch_import(request: Request, payload: dict):
    uid = _uid(request)
    inserted, skipped = 0, []
    for table, cols, items in (
        ("awards", AWARD_COLS, payload.get("awards", [])),
        ("positions", POSITION_COLS, payload.get("positions", [])),
    ):
        date_col = _DEDUP_DATE_COL[table]
        existing = {
            (str(r.get("title") or "").strip(), str(r.get(date_col) or "").strip())
            for r in db.get_rows_where(table, "user_id=?", (uid,))
        }
        for it in items:
            if not isinstance(it, dict):
                continue
            data = _pick(it, cols)
            # 防御：标题为空（含 AI 返回 null）的条目直接丢弃，避免库里出现 NULL/"None"
            if not str(data.get("title") or "").strip():
                continue
            if "ongoing" in data:
                data["ongoing"] = _to_flag(data["ongoing"])
            key = (str(data.get("title")).strip(), str(data.get(date_col) or "").strip())
            if key in existing:
                skipped.append(key[0])
                continue
            existing.add(key)
            data["user_id"] = uid
            data["sort_order"] = db.next_sort_order(table)
            db.insert_row(table, data)
            inserted += 1
    return {"ok": True, "count": inserted, "skipped": skipped}


@app.put("/api/items/{item_type}/{item_id}")
def update_item(request: Request, item_type: str, item_id: int, payload: dict):
    uid = _uid(request)
    table, cols = ITEM_TABLES.get(item_type, (None, None))
    if not table:
        raise HTTPException(400, "item_type 必须是 award 或 position")
    current = db.get_row(table, item_id)
    if not current or current["user_id"] != uid:
        raise HTTPException(404, "条目不存在")
    incoming = _pick(payload, cols)
    # 未提交的字段沿用当前值，避免部分更新把其他字段清空
    merged = {c: (incoming[c] if c in payload else current.get(c, "")) for c in cols}
    if "ongoing" in merged:
        merged["ongoing"] = _to_flag(merged["ongoing"])
    db.update_row(table, item_id, merged)
    return {"ok": True}


@app.delete("/api/items/{item_type}/{item_id}")
def delete_item(request: Request, item_type: str, item_id: int):
    uid = _uid(request)
    table, cols = ITEM_TABLES.get(item_type, (None, None))
    if not table:
        raise HTTPException(400, "item_type 必须是 award 或 position")
    current = db.get_row(table, item_id)
    if not current or current["user_id"] != uid:
        raise HTTPException(404, "条目不存在")
    # 级联清理佐证文件：DB 记录和磁盘文件一起删，避免孤儿堆积
    for att in db.attachments_for(item_type, item_id):
        _remove_local_file(att.get("stored_path"))
        db.delete_row("item_files", att["id"])
    db.delete_row(table, item_id)
    return {"ok": True}


# ---------- 佐证文件 ----------

@app.post("/api/items/{item_type}/{item_id}/attachments")
def add_attachment(request: Request, item_type: str, item_id: int, file: UploadFile = File(...)):
    uid = _uid(request)
    table, _ = ITEM_TABLES.get(item_type, (None, None))
    parent = db.get_row(table, item_id) if table else None
    if not parent or parent["user_id"] != uid:
        raise HTTPException(404, "条目不存在")
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ATTACH_EXTS:
        raise HTTPException(400, "佐证文件格式不支持")
    name = f"{uuid.uuid4().hex}{ext}"
    path = ATTACH_DIR / name
    with open(path, "wb") as f:
        f.write(file.file.read())
    att_id = db.insert_row("item_files", {
        "user_id": uid,
        "item_type": item_type,
        "item_id": item_id,
        "filename": file.filename or "",
        "stored_path": str(path),
        "created_at": _now(),
    })
    return {"id": att_id, "url": f"/uploads/attachments/{name}", "filename": file.filename}


@app.get("/api/attachments")
def list_attachments(request: Request, item_type: str, item_id: int):
    uid = _uid(request)
    table, _ = ITEM_TABLES.get(item_type, (None, None))
    parent = db.get_row(table, item_id) if table else None
    if not parent or parent["user_id"] != uid:
        raise HTTPException(404, "条目不存在")
    rows = db.attachments_for(item_type, item_id)
    for r in rows:
        r["url"] = "/uploads/attachments/" + os.path.basename(r.get("stored_path") or "")
    return rows


@app.delete("/api/attachments/{att_id}")
def delete_attachment(request: Request, att_id: int):
    uid = _uid(request)
    row = db.get_row("item_files", att_id)
    if not row or row["user_id"] != uid:
        raise HTTPException(404, "佐证文件不存在")
    _remove_local_file(row["stored_path"])
    db.delete_row("item_files", att_id)
    return {"ok": True}


# ---------- 证书归档 ----------

CERT_COLS = ["title", "category", "issuer", "date", "level", "notes"]


def _cert_public(row):
    row = dict(row)
    row["url"] = "/uploads/certificates/" + os.path.basename(row.get("stored_path") or "")
    thumb = row.get("thumb_path") or ""
    row["thumb_url"] = "/uploads/certificates/" + os.path.basename(thumb) if thumb else ""
    return row


@app.get("/certs", response_class=HTMLResponse)
def certs_page(request: Request):
    uid = _uid(request)
    rows = [_cert_public(r) for r in db.get_rows_where("certificates", "user_id=?", (uid,), "id DESC")]
    cat_counts, year_set = {}, set()
    for r in rows:
        m = re.search(r"(20\d{2})", str(r.get("date") or ""))
        r["year"] = m.group(1) if m else "未标年份"
        year_set.add(r["year"])
        cat = r.get("category") or "其他"
        cat_counts[cat] = cat_counts.get(cat, 0) + 1
    rows.sort(key=lambda r: (_date_key(r.get("date")), -r["id"]), reverse=True)
    groups = [{"year": y,
               "items": [r for r in rows if r["year"] == y]}
              for y in sorted(year_set, key=lambda v: (v == "未标年份", v), reverse=True)]
    return templates.TemplateResponse(request, "certs.html", _base_ctx(request,
        active="certs",
        certs=rows,
        groups=groups,
        total=len(rows),
        cat_counts=dict(sorted(cat_counts.items(), key=lambda kv: -kv[1])),
        has_ai=bool(db.get_setting("api_key")),
    ))


@app.post("/api/certs/upload")
def cert_upload(request: Request, file: UploadFile = File(...)):
    """拍照/图片/PDF 上传即归档：保存 → OCR → 自动识别字段 → 生成缩略图。"""
    uid = _uid(request)
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in CERT_EXTS:
        raise HTTPException(400, f"不支持的格式：{ext or '未知'}，证书归档支持 jpg/png/webp/pdf")
    name = f"{uuid.uuid4().hex}{ext}"
    path = CERT_DIR / name
    with open(path, "wb") as f:
        f.write(file.file.read())
    raw_text = parsers.extract_text(str(path), ext)
    info = cert_organizer.analyze(raw_text, file.filename or "")
    cid = db.insert_row("certificates", {
        "user_id": uid,
        **info,
        "notes": "",
        "stored_path": str(path),
        "thumb_path": cert_organizer.make_thumbnail(str(path)),
        "ocr_text": raw_text,
        "source": file.filename or "",
        "created_at": _now(),
    })
    return _cert_public(db.get_row("certificates", cid))


@app.get("/api/certs")
def list_certs(request: Request):
    uid = _uid(request)
    rows = db.get_rows_where("certificates", "user_id=?", (uid,), "id DESC")
    for r in rows:
        r["url"] = "/uploads/certificates/" + os.path.basename(r.get("stored_path") or "")
    return rows


@app.put("/api/certs/{cert_id}")
def update_cert(request: Request, cert_id: int, payload: dict):
    uid = _uid(request)
    current = db.get_row("certificates", cert_id)
    if not current or current["user_id"] != uid:
        raise HTTPException(404, "证书不存在")
    incoming = _pick(payload, CERT_COLS)
    merged = {c: (incoming[c] if c in payload else current.get(c, "")) for c in CERT_COLS}
    db.update_row("certificates", cert_id, merged)
    return {"ok": True}


@app.delete("/api/certs/{cert_id}")
def delete_cert(request: Request, cert_id: int):
    uid = _uid(request)
    row = db.get_row("certificates", cert_id)
    if not row or row["user_id"] != uid:
        raise HTTPException(404, "证书不存在")
    _remove_local_file(row.get("stored_path"))
    _remove_local_file(row.get("thumb_path"))
    db.delete_row("certificates", cert_id)
    return {"ok": True}


@app.post("/api/certs/{cert_id}/reanalyze")
def reanalyze_cert(request: Request, cert_id: int):
    """重新 OCR + 自动整理（装好 OCR 或旋转后可再跑一次）。"""
    uid = _uid(request)
    row = db.get_row("certificates", cert_id)
    if not row or row["user_id"] != uid:
        raise HTTPException(404, "证书不存在")
    path = row.get("stored_path") or ""
    ext = os.path.splitext(path)[1].lower()
    if not path or not os.path.exists(path):
        raise HTTPException(404, "原始文件已丢失，无法重新识别")
    raw_text = parsers.extract_text(path, ext)
    info = cert_organizer.analyze(raw_text, row.get("source") or "")
    # 重新识别会覆盖标题/类别/日期等字段（前端有二次确认），仅备注始终保留
    db.update_row("certificates", cert_id, {**info, "ocr_text": raw_text})
    return _cert_public(db.get_row("certificates", cert_id))


@app.post("/api/certs/{cert_id}/to-award")
def cert_to_award(request: Request, cert_id: int):
    """把归档的证书一键转为「获奖情况」条目，并把图片复制为佐证文件。"""
    uid = _uid(request)
    cert = db.get_row("certificates", cert_id)
    if not cert or cert["user_id"] != uid:
        raise HTTPException(404, "证书不存在")
    title = (cert.get("title") or "").strip() or "未命名证书"
    date = (cert.get("date") or "").strip()
    existing = [r for r in db.get_rows_where("awards", "user_id=?", (uid,))
                if (r.get("title") or "").strip() == title and (r.get("date") or "").strip() == date]
    if existing:
        award_id, created = existing[0]["id"], False
    else:
        award_id = db.insert_row("awards", {
            "title": title,
            "date": date,
            "level": cert.get("level", ""),
            "category": cert.get("category", ""),
            "organizer": cert.get("issuer", ""),
            "description": cert.get("notes", ""),
            "source_file": "证书归档",
            "user_id": uid,
            "sort_order": db.next_sort_order("awards"),
        })
        created = True
    attached = False
    src = cert.get("stored_path") or ""
    if src and os.path.exists(src):
        dest = ATTACH_DIR / f"{uuid.uuid4().hex}{os.path.splitext(src)[1].lower()}"
        shutil.copy2(src, dest)
        db.insert_row("item_files", {
            "user_id": uid,
            "item_type": "award",
            "item_id": award_id,
            "filename": f"{title}{os.path.splitext(src)[1].lower()}",
            "stored_path": str(dest),
            "created_at": _now(),
        })
        attached = True
    msg = "已在获奖条目上补充佐证" if not created else "已创建获奖条目"
    return {"ok": True, "award_id": award_id, "created": created, "attached": attached, "message": msg}


# ---------- 个人资料 ----------

@app.get("/api/profile")
def get_profile(request: Request):
    uid = _uid(request)
    profile = db.get_row("profile", uid) or {}
    profile["skills"] = db.json_list(profile.get("skills"))
    profile["languages"] = db.json_list(profile.get("languages"))
    return {
        "profile": profile,
        "education": db.get_rows_where("education", "user_id=?", (uid,)),
        "papers": db.get_rows_where("papers", "user_id=?", (uid,)),
        "projects": db.get_rows_where("projects", "user_id=?", (uid,)),
    }


@app.put("/api/profile")
def update_profile(request: Request, payload: dict):
    uid = _uid(request)
    data = payload.get("profile", payload)
    current = db.get_row("profile", uid) or {}
    merged = {}
    for c in PROFILE_COLS:
        if c in ("skills", "languages"):
            # 只有请求里显式带了才更新；否则沿用当前值，避免误清空
            if c in data:
                v = data.get(c)
                merged[c] = v if isinstance(v, str) else json.dumps(v or [], ensure_ascii=False)
            else:
                merged[c] = current.get(c) or "[]"
        elif c in data:
            merged[c] = "" if data.get(c) is None else data[c]
        else:
            merged[c] = current.get(c, "")
    db.update_row("profile", uid, _pick(merged, PROFILE_COLS))

    result = {"ok": True}
    for table, cols in (("education", EDU_COLS), ("papers", PAPER_COLS), ("projects", PROJECT_COLS)):
        if table not in payload:
            # 未提交该分组：完全不动，避免误清空
            continue
        # 返回与提交行对齐的 id 列表，前端回填后可防止自动保存重复插入
        result[table] = _sync_rows(table, cols, payload.get(table) or [], uid)
    return result


def _sync_rows(table, cols, rows, uid):
    """按 id 同步：已有行更新、新行插入、仅删除用户删掉的旧行。
    返回与 rows 对齐的 id 列表（新行为新 id，失效行对应位置为 None）。"""
    existing_ids = {r["id"] for r in db.get_rows_where(table, "user_id=?", (uid,))}
    seen = set()
    out_ids = []
    for i, r in enumerate(rows):
        if not isinstance(r, dict):
            out_ids.append(None)
            continue
        d = _pick(r, cols)
        if "ongoing" in d:
            d["ongoing"] = _to_flag(d["ongoing"])
        d["user_id"] = uid
        d["sort_order"] = i
        rid = r.get("id")
        try:
            rid = int(rid) if rid not in (None, "") else None
        except (TypeError, ValueError):
            rid = None
        if rid is None:
            new_id = db.insert_row(table, d)
            seen.add(new_id)
            out_ids.append(new_id)
        elif rid in existing_ids:
            db.update_row(table, rid, d)
            seen.add(rid)
            out_ids.append(rid)
        else:
            # 带 id 但库里已不存在：忽略不重建
            out_ids.append(None)
    for rid in existing_ids - seen:
        db.delete_row(table, rid)
    return out_ids


@app.put("/api/profile/basic")
def update_basic(request: Request, payload: dict):
    """只更新基本信息字段，不影响教育/论文/项目等其他数据。"""
    db.update_row("profile", _uid(request), _pick(payload, ["name", "gender", "birth_date", "phone", "email", "address", "hometown"]))
    return {"ok": True}


@app.post("/api/profile/summary")
def set_summary(request: Request, payload: dict):
    db.update_row("profile", _uid(request), {"summary": payload.get("summary", "")})
    return {"ok": True}


@app.post("/api/assistant/avatar")
def upload_assistant_avatar(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in {".jpg", ".jpeg", ".png", ".webp"}:
        raise HTTPException(400, "头像只支持 jpg/png/webp")
    name = f"assistant_{uuid.uuid4().hex}{ext}"
    path = PHOTO_DIR / name
    with open(path, "wb") as f:
        f.write(file.file.read())
    old = db.get_setting("assistant_avatar")
    db.set_setting("assistant_avatar", f"/uploads/photos/{name}")
    if old != f"/uploads/photos/{name}":
        _remove_upload_file(old)
    return {"ok": True, "url": f"/uploads/photos/{name}"}


@app.post("/api/profile/photo")
def upload_photo(request: Request, file: UploadFile = File(...)):
    uid = _uid(request)
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in {".jpg", ".jpeg", ".png", ".webp"}:
        raise HTTPException(400, "照片只支持 jpg/png/webp")
    name = f"photo_{uuid.uuid4().hex}{ext}"
    path = PHOTO_DIR / name
    with open(path, "wb") as f:
        f.write(file.file.read())
    current = db.get_row("profile", uid) or {}
    new_url = f"/uploads/photos/{name}"
    db.update_row("profile", uid, {"photo_path": new_url})
    if current.get("photo_path") != new_url:
        _remove_upload_file(current.get("photo_path"))
    return {"ok": True, "url": new_url}


def _clear_table(table):
    conn = db.get_conn()
    conn.execute(f"DELETE FROM {table}")
    conn.commit()
    conn.close()


# ---------- 账户管理 ----------

@app.post("/api/users")
def create_user(response: Response, payload: dict):
    name = (payload.get("name") or "").strip()[:20]
    if not name:
        raise HTTPException(400, "账户名称不能为空")
    uid = db.insert_row("users", {"name": name, "created_at": _now()})
    db.insert_profile(uid)
    # 创建后立即切换到新账户
    response.set_cookie("uid", str(uid), max_age=3600 * 24 * 365, samesite="lax")
    return {"ok": True, "id": uid}


@app.put("/api/users/{uid}")
def rename_user(uid: int, payload: dict):
    if not db.get_row("users", uid):
        raise HTTPException(404, "账户不存在")
    name = (payload.get("name") or "").strip()[:20]
    if not name:
        raise HTTPException(400, "账户名称不能为空")
    db.update_row("users", uid, {"name": name})
    return {"ok": True}


@app.post("/api/users/{uid}/switch")
def switch_user(uid: int, response: Response):
    if not db.get_row("users", uid):
        raise HTTPException(404, "账户不存在")
    response.set_cookie("uid", str(uid), max_age=3600 * 24 * 365, samesite="lax")
    return {"ok": True}


@app.delete("/api/users/{uid}")
def delete_user(uid: int, response: Response):
    users = db.all_rows("users", "id ASC")
    if len(users) <= 1:
        raise HTTPException(400, "至少保留一个账户，不能删除")
    if not db.get_row("users", uid):
        raise HTTPException(404, "账户不存在")
    # 级联删除该账户的全部数据与磁盘文件
    for table in ("awards", "positions", "education", "papers", "projects"):
        db.delete_rows_where(table, "user_id=?", (uid,))
    for f in db.get_rows_where("files", "user_id=?", (uid,), "id"):
        _remove_local_file(f["stored_path"])
    db.delete_rows_where("files", "user_id=?", (uid,))
    for att in db.get_rows_where("item_files", "user_id=?", (uid,), "id"):
        _remove_local_file(att.get("stored_path"))
    db.delete_rows_where("item_files", "user_id=?", (uid,))
    for c in db.get_rows_where("certificates", "user_id=?", (uid,), "id"):
        _remove_local_file(c.get("stored_path"))
        _remove_local_file(c.get("thumb_path"))
    db.delete_rows_where("certificates", "user_id=?", (uid,))
    db.set_setting(f"mobile_token:{uid}", "")
    db.delete_row("profile", uid)
    db.delete_row("users", uid)
    remaining = db.all_rows("users", "id ASC")
    response.delete_cookie("uid")
    if remaining:
        response.set_cookie("uid", str(remaining[0]["id"]), max_age=3600 * 24 * 365, samesite="lax")
    return {"ok": True}


# ---------- 手机扫码直传 ----------

def _lan_ip():
    """探测本机局域网 IP；失败时回环地址（扫码不可用但不报错）。"""
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(("8.8.8.8", 80))
        return s.getsockname()[0]
    except Exception:
        return "127.0.0.1"
    finally:
        s.close()


def _mobile_token(uid: int, force: bool = False):
    key = f"mobile_token:{uid}"
    tok = db.get_setting(key)
    if force or not tok:
        tok = uuid.uuid4().hex
        db.set_setting(key, tok)
    return tok


@app.get("/api/mobile/link")
def mobile_link(request: Request):
    uid = _uid(request)
    tok = _mobile_token(uid)
    ip = _lan_ip()
    port = request.url.port or 8000
    return {
        "url": f"http://{ip}:{port}/m?t={tok}",
        "token": tok,
        "ip": ip,
        "lan": os.environ.get("RESUME_LAN") == "1",
        "pair": _pair_code() if os.environ.get("RESUME_LAN") == "1" else "",
    }


@app.post("/api/mobile/rotate-token")
def mobile_rotate(request: Request):
    return {"ok": True, "token": _mobile_token(_uid(request), force=True)}


@app.get("/pair", response_class=HTMLResponse)
def pair_page(request: Request, next: str = "/"):
    if not next.startswith("/") or next.startswith("//"):
        next = "/"
    return templates.TemplateResponse(request, "pair.html", {"next": next})


@app.post("/api/pair")
def pair_submit(response: Response, payload: dict):
    code = (payload.get("code") or "").strip().upper()
    expect = _pair_code()
    if not secrets.compare_digest(code, expect):
        raise HTTPException(400, "配对码不正确")
    response.set_cookie("pair", expect, max_age=3600 * 24 * 30, samesite="lax")
    nxt = payload.get("next") or "/"
    if not nxt.startswith("/") or nxt.startswith("//"):
        nxt = "/"
    return {"ok": True, "next": nxt}


@app.post("/api/auth/logout")
def auth_logout(response: Response, payload: dict):
    """退出当前设备：scope=pair 清除局域网配对；scope=uid 退回默认账户；all 两者都清。"""
    scope = payload.get("scope") or "uid"
    if scope in ("pair", "all"):
        response.delete_cookie("pair")
    if scope in ("uid", "all"):
        response.delete_cookie("uid")
    return {"ok": True}


@app.get("/qr.svg")
def qr_svg(d: str = ""):
    import qrcode
    import qrcode.image.svg
    if not d:
        raise HTTPException(400, "缺少二维码内容")
    if len(d) > 512:
        raise HTTPException(400, "二维码内容过长")
    img = qrcode.make(d, image_factory=qrcode.image.svg.SvgPathImage, box_size=14, border=2)
    buf = io.BytesIO()
    img.save(buf)
    return Response(content=buf.getvalue(), media_type="image/svg+xml")


@app.get("/m", response_class=HTMLResponse)
def mobile_page(request: Request, t: str = ""):
    uid = None
    if t:
        for u in db.all_rows("users", "id ASC"):
            if t and t == db.get_setting(f"mobile_token:{u['id']}"):
                uid = u["id"]
                break
    if uid is None:
        return HTMLResponse(
            "<meta charset='utf-8'><body style='font-family:sans-serif;text-align:center;padding-top:80px;color:#555'>"
            "<h3>链接无效或已被重置</h3><p>请在电脑端首页重新获取二维码</p></body>", status_code=404)
    user = db.get_row("users", uid) or {"name": "?"}
    return templates.TemplateResponse(request, "mobile_upload.html", {
        "uid": uid, "user_name": user["name"], "token": t,
    })


if os.environ.get("RESUME_LAN") == "1":
    print(f"[LAN] 配对码: {_pair_code()}  （手机等设备首次打开页面时需输入，每次启动刷新）", flush=True)


# ---------- 首次启动向导 / 示例数据 ----------

@app.get("/setup", response_class=HTMLResponse)
def setup_page(request: Request):
    """初始化向导；已初始化的实例也能随时回来修改账户名/头像/AI 配置。"""
    uid = _uid(request)
    user = db.get_row("users", uid) or {}
    profile = db.get_row("profile", uid) or {}
    return templates.TemplateResponse(request, "setup.html", {
        "initialized": db.get_setting("initialized") == "1",
        "account_name": profile.get("name") or user.get("name") or "",
        "api_base": db.get_setting("api_base", ""),
        "ai_model": db.get_setting("model", ""),
        "has_key": bool(db.get_setting("api_key")),
        "assistant_avatar": db.get_setting("assistant_avatar") or "",
    })


def _friendly_conn_err(e) -> str:
    s = str(e)
    if "timed out" in s or "timeout" in s.lower():
        return "连接超时：请检查网络；OpenAI 等境外服务可能需要代理"
    if "refused" in s.lower():
        return "连接被拒绝：服务未运行或地址/端口不对（本地 Ollama 需先执行 ollama serve）"
    if "Name or service" in s or "getaddrinfo" in s or "No such host" in s:
        return "域名解析失败：请检查接口地址拼写"
    return s[:200]


@app.post("/api/ai/test")
def ai_test(payload: dict):
    """验证 OpenAI 兼容接口连通性：优先 GET /models（8 秒超时）；
    仅当服务端明确表示没有该端点（404/405）时才退回一次最小对话请求。
    其他失败（超时/拒绝/鉴权）立即返回人话错误，不让页面干等。"""
    base = (payload.get("api_base") or "").strip().rstrip("/")
    key = (payload.get("api_key") or "").strip()
    model = (payload.get("model") or "").strip()
    if not base or not model:
        return {"ok": False, "error": "接口地址与模型名称不能为空"}
    import urllib.request
    import urllib.error
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {key}"}
    try:
        req = urllib.request.Request(base + "/models", headers=headers)
        with urllib.request.urlopen(req, timeout=8) as r:
            r.read()
        return {"ok": True}
    except urllib.error.HTTPError as e:
        if e.code in (404, 405):
            pass  # 该兼容端点不提供 /models，转而试一次最小对话
        elif e.code in (401, 403):
            return {"ok": False, "error": f"HTTP {e.code}：API Key 无效或没有权限"}
        else:
            return {"ok": False, "error": f"HTTP {e.code}：{e.reason}"}
    except Exception as e:
        return {"ok": False, "error": _friendly_conn_err(e)}
    try:
        body = json.dumps({"model": model,
                           "messages": [{"role": "user", "content": "ping"}],
                           "max_tokens": 1}).encode()
        req = urllib.request.Request(base + "/chat/completions", data=body, headers=headers)
        with urllib.request.urlopen(req, timeout=12) as r:
            r.read()
        return {"ok": True}
    except urllib.error.HTTPError as e:
        return {"ok": False, "error": f"接口可达但返回 HTTP {e.code}：请检查模型名称与 API Key"}
    except Exception as e:
        return {"ok": False, "error": _friendly_conn_err(e)}


_DEMO_PROFILE = {
    "name": "李明", "gender": "男", "birth_date": "2003年5月",
    "phone": "138****0000", "email": "liming@example.com", "address": "某省某市",
    "summary": "示例大学计算机科学与技术专业本科生，专业排名前 5%。曾获全国大学生数学建模竞赛省级一等奖、"
               "校级一等奖学金；担任班长期间组织多项班级活动。熟悉 Python 与数据分析，具备扎实的编程基础和团队协作能力。",
}
_DEMO_EDU = [{"school": "示例大学", "degree": "本科", "major": "计算机科学与技术",
              "start": "2022年9月", "end": "2026年6月", "notes": "GPA 3.8/4.0，专业排名前 5%", "ongoing": 0}]
_DEMO_AWARDS = [
    {"title": "全国大学生数学建模竞赛省级一等奖", "level": "省级", "category": "竞赛",
     "date": "2024年11月", "organizer": "竞赛组委会"},
    {"title": "校级一等奖学金", "level": "校级", "category": "学业",
     "date": "2024年10月", "organizer": "示例大学"},
    {"title": "校三好学生", "level": "校级", "category": "荣誉",
     "date": "2023年12月", "organizer": "示例大学"},
]
_DEMO_POSITIONS = [
    {"title": "班长", "org": "计算机2201班", "start": "2023年9月", "end": "",
     "ongoing": 1, "description": "负责班级日常管理与活动组织，班级获评校级优秀班集体。"},
    {"title": "学习部干事", "org": "校学生会", "start": "2022年10月", "end": "2023年6月",
     "ongoing": 0, "description": "协助组织学习经验交流会与学风建设月活动。"},
]
_DEMO_PAPERS = [{"title": "基于示例数据的学习行为分析方法", "authors": "李明，张三",
                 "venue": "示例学报", "year": "2025", "volume": "42", "issue": "3",
                 "pages": "15-20", "paper_type": "普通期刊"}]
_DEMO_PROJECTS = [{"name": "校园二手交易小程序", "role": "项目负责人",
                   "start": "2024年3月", "end": "2024年9月", "ongoing": 0,
                   "description": "带领 4 人团队完成需求分析、后端开发与上线运营，累计注册用户 1200+。"}]
_DEMO_SKILLS = ["Python", "SQL", "Office 办公套件"]
_DEMO_LANGS = ["英语（CET-6）"]


@app.post("/api/setup/finish")
def setup_finish(response: Response, payload: dict):
    name = (payload.get("account_name") or "").strip()[:20]
    users = db.all_rows("users", "id ASC")
    if name and users:
        db.update_row("users", users[0]["id"], {"name": name})
        db.update_row("profile", users[0]["id"], {"name": name})
    if payload.get("load_demo"):
        uid = users[0]["id"] if users else 1
        empty = all(not db.get_rows_where(t, "user_id=?", (uid,))
                    for t in ("awards", "positions", "education", "papers", "projects"))
        if empty:
            seed_demo_data(uid)
    if not payload.get("skip_ai"):
        key = (payload.get("api_key") or "").strip()
        if key:
            db.set_setting("api_key", key)
            db.set_setting("api_base", (payload.get("api_base") or "https://api.deepseek.com/v1").strip())
            db.set_setting("model", (payload.get("model") or "deepseek-chat").strip())
    db.set_setting("initialized", "1")
    if users:
        response.set_cookie("uid", str(users[0]["id"]), max_age=3600 * 24 * 365, samesite="lax")
    return {"ok": True}


def seed_demo_data(uid: int):
    """载入虚构人物「李明」的示例资料（仅在空白账户上调用）。"""
    prof = {k: v for k, v in _DEMO_PROFILE.items()}
    db.update_row("profile", uid, _pick(prof, PROFILE_COLS))
    for table, cols, items in (
        ("education", EDU_COLS, _DEMO_EDU),
        ("awards", AWARD_COLS, _DEMO_AWARDS),
        ("positions", POSITION_COLS, _DEMO_POSITIONS),
        ("papers", PAPER_COLS, _DEMO_PAPERS),
        ("projects", PROJECT_COLS, _DEMO_PROJECTS),
    ):
        for i, it in enumerate(items):
            data = _pick(it, cols)
            data["user_id"] = uid
            data["sort_order"] = i
            db.insert_row(table, data)
    db.update_row("profile", uid, {"skills": json.dumps(_DEMO_SKILLS, ensure_ascii=False),
                                   "languages": json.dumps(_DEMO_LANGS, ensure_ascii=False)})


@app.post("/api/demo/load")
def demo_load(request: Request):
    uid = _uid(request)
    empty = all(not db.get_rows_where(t, "user_id=?", (uid,))
                for t in ("awards", "positions", "education", "papers", "projects"))
    if not empty:
        raise HTTPException(400, "当前账户已有数据，仅空白账户可载入示例数据")
    seed_demo_data(uid)
    return {"ok": True}


@app.post("/api/demo/clear")
def demo_clear(request: Request):
    """清空当前账户的全部简历数据（保留账户名），用于移除示例或重新开始。"""
    uid = _uid(request)
    for table in ("awards", "positions", "education", "papers", "projects"):
        db.delete_rows_where(table, "user_id=?", (uid,))
    blank = {c: "" for c in PROFILE_COLS}
    blank["skills"] = "[]"
    blank["languages"] = "[]"
    photo = db.get_row("profile", uid) or {}
    keep_photo = photo.get("photo_path") or ""
    db.update_row("profile", uid, blank)
    if keep_photo:
        _remove_upload_file(keep_photo)
    return {"ok": True}


# ---------- 备份与恢复 ----------

@app.get("/api/backups")
def list_backups():
    """列出 data_backups/ 下的全部快照（新→旧）。"""
    bdir = Path(db.BACKUP_DIR)
    out = []
    if bdir.exists():
        for f in sorted(bdir.glob("data_*.db"), reverse=True):
            try:
                created = datetime.strptime(f.stem[5:], "%Y%m%d_%H%M%S").strftime("%Y-%m-%d %H:%M")
            except ValueError:
                created = ""
            out.append({"filename": f.name, "size": f.stat().st_size, "created_at": created})
    return out


@app.post("/api/backups/restore")
def restore_backup(payload: dict):
    """把指定快照恢复为当前数据库；恢复前先把现状再快照一份兜底。"""
    name = str(payload.get("filename") or "")
    if not re.fullmatch(r"data_\d{8}_\d{6}\.db", name):
        raise HTTPException(400, "备份文件名不合法")
    src = Path(db.BACKUP_DIR) / name
    if not src.is_file():
        raise HTTPException(404, "备份不存在")
    try:
        db.backup_db()
    except Exception:
        pass
    s = sqlite3.connect(str(src))
    d = sqlite3.connect(db.DB_PATH)
    try:
        s.backup(d)
    finally:
        d.close()
        s.close()
    return {"ok": True}


# ---------- 模板管理 ----------

@app.get("/api/templates")
def list_templates():
    return db.all_rows("resume_templates", "builtin DESC, id ASC")


@app.post("/api/templates")
def create_template(payload: dict):
    name = (payload.get("name") or "").strip()
    html = payload.get("html") or ""
    if not name or not html.strip():
        raise HTTPException(400, "模板名称和内容不能为空")
    if db.get_row_by("resume_templates", "name", name):
        raise HTTPException(400, f"模板 {name} 已存在")
    db.insert_row("resume_templates", {
        "name": name,
        "title": name,
        "html": html,
        "builtin": 0,
        "created_at": _now(),
    })
    return {"ok": True}


@app.delete("/api/templates/{name}")
def delete_template(name: str):
    row = db.get_row_by("resume_templates", "name", name)
    if not row:
        raise HTTPException(404, "模板不存在")
    if row["builtin"]:
        raise HTTPException(400, "内置模板不能删除")
    db.delete_row("resume_templates", row["id"])
    return {"ok": True}


# ---------- 设置 / AI ----------

def _mask_secret(key: str) -> str:
    if not key:
        return ""
    if len(key) <= 8:
        return key[:2] + "***"
    return f"{key[:3]}***{key[-4:]}"


@app.get("/api/settings")
def get_settings():
    return {
        "api_key": _mask_secret(db.get_setting("api_key")),
        "api_base": db.get_setting("api_base", "https://api.deepseek.com/v1"),
        "model": db.get_setting("model", "deepseek-chat"),
        "assistant_avatar": db.get_setting("assistant_avatar"),
    }


@app.put("/api/settings")
def update_settings(payload: dict):
    key = payload.get("api_key")
    # 空 / 未提交 / 掩码值（含 ***）都视为「不修改」；输入新 Key 即覆盖
    if key and "***" not in key:
        db.set_setting("api_key", str(key).strip())
    if payload.get("api_base"):
        db.set_setting("api_base", payload["api_base"])
    if payload.get("model"):
        db.set_setting("model", payload["model"])
    return {"ok": True}


def _llm_chat(messages, temperature=0.5):
    api_key = db.get_setting("api_key")
    if not api_key:
        raise HTTPException(400, "还没有配置 AI 的 API Key，请先在首页填写")
    url = db.get_setting("api_base", "https://api.deepseek.com/v1").rstrip("/") + "/chat/completions"
    body = {
        "model": db.get_setting("model", "deepseek-chat"),
        "temperature": temperature,
        "messages": messages,
    }
    import urllib.request
    req = urllib.request.Request(
        url,
        data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
    )
    try:
        with urllib.request.urlopen(req, timeout=180) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(502, f"AI 调用失败：{e}")
    return data["choices"][0]["message"]["content"]


@app.post("/api/ai/extract")
def ai_extract(request: Request, payload: dict):
    uid = _uid(request)
    file_id = payload.get("file_id")
    row = db.get_row("files", file_id) if file_id else None
    if not row or row["user_id"] != uid:
        raise HTTPException(404, "文件不存在")
    if not db.get_setting("api_key"):
        raise HTTPException(400, "还没有配置 AI 的 API Key，请先在首页填写")
    try:
        result = extractor.ai_extract(
            row["raw_text"],
            db.get_setting("api_base", "https://api.deepseek.com/v1"),
            db.get_setting("api_key"),
            db.get_setting("model", "deepseek-chat"),
        )
    except Exception as e:
        raise HTTPException(502, f"AI 调用失败：{e}")
    return result


def _resume_data_block(uid: int):
    ctx = resume_context(uid)
    p = ctx["profile"]
    lines = [f"姓名：{p.get('name') or '（未填）'}", f"联系方式：{p.get('phone') or ''} {p.get('email') or ''}"]
    lines.append(f"个人简介（当前）：{p.get('summary') or '（未填）'}")
    if ctx["education"]:
        lines.append("教育背景：" + "；".join(
            f"{e.get('school')} {e.get('degree')} {e.get('major')} ({e.get('start')}-{e.get('end')})"
            for e in ctx["education"]))
    if ctx["papers"]:
        lines.append("论文：" + "；".join(
            f"{a.get('title')}《{a.get('venue')}》{a.get('year')}" for a in ctx["papers"]))
    if ctx["projects"]:
        lines.append("项目：" + "；".join(
            f"{pr.get('name')}({pr.get('role')})" for pr in ctx["projects"]))
    if ctx["awards"]:
        lines.append("获奖：" + "；".join(
            f"{a.get('title')}({a.get('level')}, {a.get('date')})" for a in ctx["awards"]))
    if ctx["positions"]:
        lines.append("任职：" + "；".join(
            f"{po.get('title')}@{po.get('org')}({po.get('start')}-{po.get('end')})" for po in ctx["positions"]))
    if p.get("skills"):
        lines.append("技能：" + "、".join(p["skills"]))
    return "\n".join(lines)


@app.post("/api/chat")
def chat(request: Request, payload: dict):
    messages = payload.get("messages") or []
    if not messages:
        raise HTTPException(400, "消息不能为空")
    system = (
        "你是「小点」，一位友好、耐心的中文简历助手，帮助用户完善个人介绍和简历。\n\n"
        "用户的简历数据如下：\n" + _resume_data_block(_uid(request)) + "\n\n"
        "你的工作方式：\n"
        "1. 先了解用户这份简历的用途（求职/保研/考研复试/奖学金评定/评优等）和目标岗位或方向；\n"
        "2. 基于用途给个人简介、获奖与任职描述、项目经历等提出具体、可落地的修改建议；\n"
        "3. 用户要求改写时，直接给出改写后的中文文本（个人简介一般 80~150 字，突出亮点和数据）；\n"
        "4. 可以指出简历里缺失的信息并提醒补充；\n"
        "5. 回答保持简洁有条理，适当用短句和换行，不要长篇大论。"
    )
    try:
        reply = _llm_chat([{"role": "system", "content": system}] + messages[-24:], temperature=0.7)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(502, f"AI 调用失败：{e}")
    return {"reply": reply}


# ---------- 导出 ----------

def _docx_style(doc):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def set_font(run, size=11, bold=False):
        run.font.name = "微软雅黑"
        run.font.size = Pt(size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return set_font


@app.get("/resume.docx")
def resume_docx(request: Request):
    from docx import Document
    from docx.shared import Pt, Inches

    set_font = _docx_style(None)
    ctx = resume_context(_uid(request))
    p = ctx["profile"]
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(0.8)

    h = doc.add_heading(level=0)
    h.alignment = 1
    set_font(h.add_run(p.get("name") or "我的简历"), 20, True)
    contact = "  |  ".join(x for x in [p.get("phone"), p.get("email"), p.get("address")] if x)
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = 1
        set_font(cp.add_run(contact), 10)

    def section(title):
        sp = doc.add_paragraph()
        set_font(sp.add_run(title), 13, True)
        sp.paragraph_format.space_before = Pt(10)

    def line(label, value):
        if value:
            para = doc.add_paragraph()
            set_font(para.add_run(f"{label}：{value}" if label else value), 10.5)

    if p.get("summary"):
        section("个人简介")
        line("", p.get("summary"))
    if ctx["education"]:
        section("教育背景")
        for e in ctx["education"]:
            line("", f"{e.get('school')} | {e.get('degree')} | {e.get('major')}  ({e.get('start')} - {e.get('end')})")
            if e.get("notes"):
                line("", e.get("notes"))
    if ctx["papers"]:
        section("论文发表")
        for a in ctx["papers"]:
            line("", a.get("citation") or paper_citation(a))
    if ctx["projects"]:
        section("项目经历")
        for pr in ctx["projects"]:
            line("", f"{pr.get('name')}（{pr.get('role')}，{pr.get('start')} - {pr.get('end')}）")
            if pr.get("description"):
                line("", pr.get("description"))
    if ctx["awards"]:
        section("获奖情况")
        for a in ctx["awards"]:
            line("", f"{a.get('date')}  {a.get('title')}（{a.get('level')}）")
            if a.get("organizer"):
                line("颁奖单位", a.get("organizer"))
    if ctx["positions"]:
        section("任职经历")
        for po in ctx["positions"]:
            line("", f"{po.get('start')} - {po.get('end')}  {po.get('title')}  {po.get('org')}")
            if po.get("description"):
                line("", po.get("description"))
    if p.get("skills") or p.get("languages"):
        section("技能与其他")
        if p.get("skills"):
            line("专业技能", "、".join(p["skills"]))
        if p.get("languages"):
            line("语言", "、".join(p["languages"]))

    buf = io.BytesIO()
    doc.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=my_resume.docx"},
    )


@app.get("/intro.docx")
def intro_docx(request: Request):
    """独立输出「个人介绍」Word 文档：照片 + 基本信息 + 个人简介。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches

    set_font = _docx_style(None)
    p = db.get_row("profile", _uid(request)) or {}
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(0.9)

    photo_url = p.get("photo_path") or ""
    if photo_url:
        local = Path(str(photo_url).replace("/uploads/photos/", str(PHOTO_DIR / "")))
        if local.exists():
            pic = doc.add_paragraph()
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.add_run().add_picture(str(local), width=Inches(1.5))

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(h.add_run(p.get("name") or "我的姓名"), 22, True)

    contact = "  |  ".join(x for x in [p.get("gender"), p.get("birth_date"), p.get("hometown"),
                                       p.get("phone"), p.get("email"), p.get("address")] if x)
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cp.add_run(contact), 10)

    sp = doc.add_paragraph()
    set_font(sp.add_run("个人介绍"), 15, True)
    sp.paragraph_format.space_before = Pt(14)
    body = doc.add_paragraph()
    set_font(body.add_run(p.get("summary") or "（个人简介尚未填写，请到「个人资料」页填写。）"), 11)

    buf = io.BytesIO()
    doc.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=my_intro.docx"},
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
