"""把各种格式的文件统一提取成纯文本。OCR 组件采用惰性加载，装不上也不影响其他功能。"""
import os
import tempfile
import threading
import uuid

_ocr_engine = None


def get_ocr():
    global _ocr_engine
    if _ocr_engine is None:
        from rapidocr_onnxruntime import RapidOCR
        _ocr_engine = RapidOCR()
    return _ocr_engine


def warmup_ocr():
    """后台线程预热 OCR 模型，避免首次上传时才加载导致长时间等待。"""
    def _run():
        try:
            get_ocr()
        except Exception:
            pass
    threading.Thread(target=_run, daemon=True).start()


IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tif", ".tiff")


def fix_exif_orientation(path):
    """按 EXIF Orientation 把手机拍的照片转正；无需旋转时原样返回路径。"""
    try:
        from PIL import Image, ImageOps
        with Image.open(path) as img:
            try:
                orientation = img.getexif().get(274)
            except Exception:
                orientation = None
            if orientation in (None, 1):
                return path
            fixed = ImageOps.exif_transpose(img)
            if fixed.mode not in ("RGB", "L"):
                fixed = fixed.convert("RGB")
            tmp = os.path.join(tempfile.gettempdir(), f"_rs_up_{os.getpid()}_{uuid.uuid4().hex}.jpg")
            fixed.save(tmp, "JPEG", quality=92)
            return tmp
    except Exception:
        return path


def rotate_image_file(path, cw_angle):
    """把图片文件顺时针旋转 cw_angle 度并覆盖原文件（用于识别效果不好时的手动纠正）。"""
    try:
        from PIL import Image
    except ImportError:
        return False
    ext = os.path.splitext(path)[1].lower()
    tmp = f"{path}.rot-{uuid.uuid4().hex}.tmp"
    try:
        with Image.open(path) as img:
            out = img.rotate(-cw_angle, expand=True)  # PIL 的正角度是逆时针
            if ext in (".jpg", ".jpeg"):
                if out.mode != "RGB":
                    out = out.convert("RGB")
                out.save(tmp, "JPEG", quality=92)
            else:
                out.save(tmp, "PNG")
        os.replace(tmp, path)
        return True
    except Exception:
        if os.path.exists(tmp):
            try:
                os.unlink(tmp)
            except OSError:
                pass
        return False


def ocr_image(image_path):
    """返回 (拼接文本, 明细列表)。OCR 不可用时返回空。"""
    try:
        engine = get_ocr()
        result, _ = engine(image_path)
    except Exception:
        return "", []
    if not result:
        return "", []
    lines = [item[1] for item in result if item and item[1]]
    return "\n".join(lines), result


def extract_docx(path):
    try:
        from docx import Document
    except ImportError:
        return ""
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(path):
    try:
        import pdfplumber
    except ImportError:
        return ""
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    joined = "\n".join(parts)
    if len(joined.strip()) < 30:
        return extract_pdf_ocr(path)
    return joined


def extract_pdf_ocr(path):
    """扫描版 PDF：把每一页渲染成图片再 OCR。临时文件名带 uuid，避免并发时互相覆盖。"""
    try:
        import pdfplumber
    except ImportError:
        return ""
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            tmp = None
            try:
                img = page.to_image(resolution=200)
                tmp = os.path.join(tempfile.gettempdir(), f"_rs_{os.getpid()}_{uuid.uuid4().hex}.png")
                img.save(tmp)
                text, _ = ocr_image(tmp)
                parts.append(text or "")
            except Exception:
                parts.append("")
            finally:
                if tmp and os.path.exists(tmp):
                    try:
                        os.unlink(tmp)
                    except OSError:
                        pass
    return "\n".join(parts)


def extract_image(path):
    tmp = fix_exif_orientation(path)
    try:
        text, _ = ocr_image(tmp)
    finally:
        if tmp != path and os.path.exists(tmp):
            try:
                os.unlink(tmp)
            except OSError:
                pass
    return text


def extract_text(path, ext):
    ext = ext.lower()
    if ext == ".docx":
        return extract_docx(path)
    if ext == ".pdf":
        return extract_pdf(path)
    if ext in (".jpg", ".jpeg", ".png", ".bmp", ".webp", ".tif", ".tiff"):
        return extract_image(path)
    return ""
